#!/usr/bin/env python3
"""
ValueMitra — Token Map Reference Generator
Reads all 21 templates (originals from backup), simulates the tokenizer,
and produces a single reference .docx with one section per template showing
Field Label → Token mappings for review.

Usage:
    pip install python-docx
    python3 scripts/generate_token_map.py
    # Output: scripts/Token_Map_Reference.docx
"""

import re, zipfile
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BACKUP_DIR    = Path(__file__).parent.parent / 'apps/api/uploads/templates_backup'
TEMPLATES_DIR = Path(__file__).parent.parent / 'apps/api/uploads/templates'
OUTPUT        = Path(__file__).parent / 'Token_Map_Reference.docx'

# ─── Helpers ──────────────────────────────────────────────────────────────────

def cell_text(xml):
    t = re.sub(r'<[^>]+>', '', xml)
    return re.sub(r'\s+', ' ', t).strip()

def normalize(s):
    return re.sub(r'\s+', ' ', s.lower()
                  .replace('/', ' ').replace('-', ' ').replace('.', ' ')
                  .replace('(', ' ').replace(')', ' ').replace(',', ' ')).strip()

def has_drawing(xml):
    return '<w:drawing>' in xml

def para_text(xml):
    t = re.sub(r'<[^>]+>', ' ', xml)
    return re.sub(r'\s+', ' ', t).strip()

# ─── Label → Token mapping (mirror of tokenize_templates.py) ─────────────────

LABEL_TOKEN_MAP = [
    ('purpose for which the valuation',     '{purposeOfValuation}'),
    ('purpose of valuation',                '{purposeOfValuation}'),
    ('date of inspection',                  '{inspectionDate}'),
    ('date on which the valuation is made', '{valuationDate}'),
    ('date on which valuation',             '{valuationDate}'),
    ('valuation date',                      '{valuationDate}'),
    ('date of valuation',                   '{valuationDate}'),
    ('name of the owner',                   '{ownerName}'),
    ('name of borrower',                    '{ownerName}'),
    ('name of applicant',                   '{ownerName}'),
    ('name of mortgagor',                   '{ownerName}'),
    ('owner name',                          '{ownerName}'),
    ('brief description of the property',   '{propertyAddress}'),
    ('brief description',                   '{propertyAddress}'),
    ('description of the property',         '{propertyAddress}'),
    ('total lease period',                  '{landTenure}'),
    ('leasehold',                           '{landTenure}'),
    ('plot no   survey no',                 '{surveyNo}'),
    ('survey no',                           '{surveyNo}'),
    ('hissa no',                            '{hissaNo}'),
    ('cts no',                              '{ctsSurveyNo}'),
    ('cts survey no',                       '{ctsSurveyNo}'),
    ('door no',                             '{municipalNo}'),
    ('plot no',                             '{municipalNo}'),
    ('municipal no',                        '{municipalNo}'),
    ('t  s  no   village',                  '{village}'),
    ('village',                             '{village}'),
    ('ward   taluka',                       '{taluka}'),
    ('taluka',                              '{taluka}'),
    ('mandal   district',                   '{district}'),
    ('district',                            '{district}'),
    ('postal address',                      '{propertyAddress}'),
    ('address of the property',             '{propertyAddress}'),
    ('location of property',                '{propertyAddress}'),
    ('city   town',                         '{district}'),
    ('pin code',                            '{pincode}'),
    ('pincode',                             '{pincode}'),
    ('landmark',                            '{landmark}'),
    ('latitude',                            '{landmark}'),
    ('high   middle   poor',                '{areaClassificationGrade}'),
    ('classification of area',              '{areaClassificationGrade}'),
    ('classification of locality',          '{areaClassificationGrade}'),
    ('urban   semi urban',                  '{urbanRuralClass}'),
    ('urban   semi-urban',                  '{urbanRuralClass}'),
    ('corporation limit',                   '{municipalBodyType}'),
    ('village panchayat',                   '{municipalBodyType}'),
    ('municipality',                        '{municipalBodyType}'),
    ('covered under any state',             '{govtEnactments}'),
    ('urban land ceiling',                  '{govtEnactments}'),
    ('extent of the site considered',       '{landAreaActualSqM}'),
    ('extent of site considered',           '{landAreaActualSqM}'),
    ('extent of the site',                  '{landAreaSqM}'),
    ('extent of site',                      '{landAreaSqM}'),
    ('size of plot',                        '{landAreaSqM}'),
    ('total extent',                        '{landAreaSqM}'),
    ('area of plot',                        '{landAreaSqM}'),
    ('land area',                           '{landAreaSqM}'),
    ('plot area',                           '{landAreaSqM}'),
    ('super built up area',                 '{superBuiltUpAreaSqFt}'),
    ('saleable area',                       '{superBuiltUpAreaSqFt}'),
    ('built up area',                       '{builtUpAreaSqFt}'),
    ('plinth area',                         '{builtUpAreaSqFt}'),
    ('carpet area',                         '{carpetAreaSqFt}'),
    ('undivided share',                     '{udsArea}'),
    ('uds',                                 '{udsArea}'),
    ('shape of land',                       '{plotShape}'),
    ('shape of plot',                       '{plotShape}'),
    ('shape',                               '{plotShape}'),
    ('type of use to which',                '{zoningClassification}'),
    ('type of use',                         '{zoningClassification}'),
    ('permissible use',                     '{permissibleUse}'),
    ('zoning',                              '{zoningClassification}'),
    ('development plan zone',               '{dpZone}'),
    ('dp zone',                             '{dpZone}'),
    ('fsi',                                 '{fsi}'),
    ('far',                                 '{fsi}'),
    ('occupied by the owner',               '{possessionStatus}'),
    ('whether occupied',                    '{possessionStatus}'),
    ('occupancy',                           '{possessionStatus}'),
    ('possession',                          '{possessionStatus}'),
    ('water potentiality',                  '{positiveFactors}'),
    ('underground sewerage',                '{positiveFactors}'),
    ('power supply',                        '{positiveFactors}'),
    ('advantage of the site',               '{positiveFactors}'),
    ('road facilities',                     '{positiveFactors}'),
    ('type of road',                        '{positiveFactors}'),
    ('width of road',                       '{positiveFactors}'),
    ('special remarks',                     '{negativeFactors}'),
    ('type of building',                    '{zoningClassification}'),
    ('type of construction',                '{structureType}'),
    ('nature of construction',              '{structureType}'),
    ('rcc   load bearing',                  '{structureType}'),
    ('year of construction',                '{yearOfConstruction}'),
    ('age of building',                     '{ageOfBuilding}'),
    ('remaining life',                      '{remainingLifeYears}'),
    ('number of floors',                    '{numberOfFloors}'),
    ('no  of floors',                       '{numberOfFloors}'),
    ('number of storeys',                   '{numberOfFloors}'),
    ('condition of the building',           '{condition}'),
    ('exterior',                            '{condition}'),
    ('date of issue and validity',          '{approvedPlanDateFormatted}'),
    ('approved plan issuing authority',     '{approvedPlanAuthority}'),
    ('approved map   plan issuing',         '{approvedPlanAuthority}'),
    ('issuing authority',                   '{approvedPlanAuthority}'),
    ('genuineness or authenticity',         '{planGenuinenessVerified}'),
    ('genuineness of approved',             '{planGenuinenessVerified}'),
    ('plan genuineness',                    '{planGenuinenessVerified}'),
    ('unauthorized construction',           '{unauthorizedConstruction}'),
    ('demolition proceedings',              '{demolitionProceedings}'),
    ('flooring',                            '{flooringType}'),
    ('wall finishing',                      '{wallFinishing}'),
    ('electrical',                          '{electricalFittings}'),
    ('sanitary',                            '{sanitaryFittings}'),
    ('ceiling height',                      '{ceilingHeightFt}'),
    ('guideline rate',                      '{govtRatePerSqM}'),
    ('circle rate',                         '{govtRatePerSqM}'),
    ('ready reckoner',                      '{govtRatePerSqM}'),
    ('jantri rate',                         '{govtRatePerSqM}'),
    ('rr rate',                             '{govtRatePerSqM}'),
    ('registrar',                           '{govtRatePerSqM}'),
    ('prevailing market rate',              '{comparablesNarrative}'),
    ('assessed   adopted rate',             '{govtRatePerSqM}'),
    ('adopted rate',                        '{govtRatePerSqM}'),
    ('estimated value of land',             '{landValue}'),
    ('value of land',                       '{landValue}'),
    ('land value',                          '{landValue}'),
    ('estimated value of building',         '{buildingValue}'),
    ('value of building',                   '{buildingValue}'),
    ('building value',                      '{buildingValue}'),
    ('fair market value',                   '{fairMarketValue}'),
    ('market value',                        '{finalValue}'),
    ('realizable value',                    '{realizableValue}'),
    ('distress value',                      '{distressSaleValue}'),
    ('distress sale value',                 '{distressSaleValue}'),
    ('forced sale value',                   '{distressSaleValue}'),
    ('final value',                         '{finalValue}'),
    ('total value',                         '{finalValue}'),
    ('value as per circle rate',            '{govtRatePerSqM}'),
    ('insurance value',                     '{insuranceValue}'),
    ('rental value',                        '{rentalValueMonthly}'),
    ('book value',                          '{bookValue}'),
    ('bank name',                           '{bankName}'),
    ('name of bank',                        '{bankName}'),
    ('bank branch',                         '{bankBranch}'),
    ('branch',                              '{bankBranch}'),
    ('loan account',                        '{loanAccountNo}'),
    ('account no',                          '{loanAccountNo}'),
    ('bank ref',                            '{bankRefNo}'),
    ('reference no',                        '{firmReferenceNo}'),
    ('report date',                         '{reportDate}'),
    ('date of report',                      '{reportDate}'),
    ('name of valuer',                      '{rvFullName}'),
    ('valuer name',                         '{rvFullName}'),
    ('ibbi reg',                            '{rvIbbiRegNo}'),
    ('registration no',                     '{rvIbbiRegNo}'),
    ('rvo',                                 '{rvRvoName}'),
    ('qualifications',                      '{rvQualifications}'),
    ('marketability',                       '{marketabilityRating}'),
    ('favourable factors',                  '{positiveFactors}'),
    ('positive factors',                    '{positiveFactors}'),
    ('negative factors',                    '{negativeFactors}'),
    ('unfavourable factors',                '{negativeFactors}'),
    ('market analysis',                     '{marketAnalysisNarrative}'),
    ('comparable transactions',             '{comparablesNarrative}'),
    ('comparables',                         '{comparablesNarrative}'),
]

BOUNDARY_LABELS = {
    'north': ('{northBoundaryDoc}', '{northBoundaryActual}'),
    'south': ('{southBoundaryDoc}', '{southBoundaryActual}'),
    'east':  ('{eastBoundaryDoc}',  '{eastBoundaryActual}'),
    'west':  ('{westBoundaryDoc}',  '{westBoundaryActual}'),
}

SECTION_PHOTO_MAP = [
    ('photograph',     'propertyPhoto'),
    ('location',       'locationMap'),
    ('ready reckoner', 'rrRateScreenshot'),
    ('rr rate',        'rrRateScreenshot'),
    ('price indicat',  'onlineRate'),
    ('online rate',    'onlineRate'),
]

def find_token(label):
    norm = normalize(label)
    for pattern, token in LABEL_TOKEN_MAP:
        if pattern in norm:
            return token
    return None

# ─── Extract mappings from a template ────────────────────────────────────────

def extract_mappings(docx_path: Path):
    """Returns list of (label, token, note) tuples for all matched cells + image paragraphs."""
    mappings = []

    with zipfile.ZipFile(docx_path, 'r') as zin:
        xml = zin.read('word/document.xml').decode('utf-8')

    # ── Table rows ────────────────────────────────────────────────────────────
    rows = re.findall(r'<w:tr[ >].*?</w:tr>', xml, re.DOTALL)
    for row in rows:
        cells = re.findall(r'<w:tc>.*?</w:tc>', row, re.DOTALL)
        if len(cells) < 2:
            continue
        texts = [cell_text(c) for c in cells]

        # Boundary rows
        first_norm = normalize(texts[0])
        matched_boundary = False
        for bkey, (doc_tok, act_tok) in BOUNDARY_LABELS.items():
            if first_norm == bkey or first_norm.startswith(bkey + ' '):
                matched_boundary = True
                if len(texts) >= 2:
                    label = texts[0] or bkey.capitalize()
                    mappings.append((label, doc_tok, 'Boundary — As per document'))
                if len(texts) >= 3:
                    mappings.append((label, act_tok, 'Boundary — Actual'))
                break
        if matched_boundary:
            continue

        # Normal rows
        label_idx = -1
        for i, t in enumerate(texts[:-1]):
            if len(t) > 8:
                label_idx = i
                break
        if label_idx < 0:
            continue

        token = find_token(texts[label_idx])
        if not token:
            continue

        # Check there IS an empty cell to fill
        has_empty = any(not texts[j] for j in range(label_idx + 1, len(texts)))
        if not has_empty:
            continue

        # Value cell position description
        num_cells = len(texts)
        mappings.append((texts[label_idx], token, f'{num_cells}-column row'))

    # ── Paragraph images (photos, maps, etc.) ─────────────────────────────────
    tbl_ranges = []
    for m in re.finditer(r'<w:tbl[ >]', xml):
        start = m.start()
        depth = 0; pos = start
        while pos < len(xml):
            if xml[pos:pos+6] == '<w:tbl' and pos+6 < len(xml) and xml[pos+6] in ' >':
                depth += 1; pos += 6
            elif xml[pos:pos+8] == '</w:tbl>':
                depth -= 1
                if depth == 0:
                    tbl_ranges.append((start, pos + 8)); break
                pos += 8
            else:
                pos += 1

    def in_table(p):
        return any(s <= p < e for s, e in tbl_ranges)

    current_section = None
    section_count = {'propertyPhoto': 0, 'onlineRate': 0}

    for m in re.finditer(r'<w:p[ >]', xml):
        if in_table(m.start()):
            continue
        start = m.start()
        end_m = re.search(r'</w:p>', xml[start:])
        if not end_m:
            continue
        end = start + end_m.end()
        p_xml = xml[start:end]
        p_txt = para_text(p_xml).lower()

        for keyword, section_key in SECTION_PHOTO_MAP:
            if keyword in p_txt:
                if section_key != current_section:
                    current_section = section_key
                    if section_key == 'propertyPhoto':
                        section_count['propertyPhoto'] = 0
                    elif section_key == 'onlineRate':
                        section_count['onlineRate'] = 0
                break

        if not has_drawing(p_xml) or not current_section:
            continue

        if current_section == 'propertyPhoto':
            section_count['propertyPhoto'] += 1
            if section_count['propertyPhoto'] > 12:
                continue
            token = f'{{%propertyPhoto{section_count["propertyPhoto"]}}}'
            label = f'Property Photo #{section_count["propertyPhoto"]}'
        elif current_section == 'locationMap':
            token = '{%locationMap}'
            label = 'Location Map'
            current_section = None
        elif current_section == 'rrRateScreenshot':
            token = '{%rrRateScreenshot}'
            label = 'RR Rate Screenshot'
            current_section = None
        elif current_section == 'onlineRate':
            section_count['onlineRate'] += 1
            if section_count['onlineRate'] > 4:
                continue
            token = f'{{%onlineRate{section_count["onlineRate"]}}}'
            label = f'Online Rate Evidence #{section_count["onlineRate"]}'
        else:
            continue

        mappings.append((label, token, 'IMAGE — paragraph-level'))

    return mappings

# ─── Styling helpers ──────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def bold_para(para, text, size=11, color=None):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))

# ─── Build the reference document ─────────────────────────────────────────────

def build_doc(all_data):
    from docx.shared import Cm as DocxCm
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = DocxCm(1.5)
        section.bottom_margin = DocxCm(1.5)
        section.left_margin   = DocxCm(1.8)
        section.right_margin  = DocxCm(1.8)

    # Title
    title = doc.add_heading('ValueMitra — Template Token Map Reference', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(16)

    intro = doc.add_paragraph(
        'This document lists every field label found in each of the 21 report templates and the {token} '
        'that the tokenizer will inject into it. Review each section carefully:\n'
        '• If a label is mapped to the WRONG token — note it with the correct token.\n'
        '• If a label that SHOULD be tokenized is missing — add it.\n'
        '• If a label should NOT be tokenized — note it.\n'
        'Return the annotated copy and the tokenizer mapping will be updated accordingly.'
    )
    intro.style.font.size = Pt(9)

    doc.add_paragraph()  # spacer

    # Legend
    legend_tbl = doc.add_table(rows=1, cols=3)
    legend_tbl.style = 'Table Grid'
    hdr = legend_tbl.rows[0].cells
    for cell, txt, color in [
        (hdr[0], '  White = Text token ({token})', 'FFFFFF'),
        (hdr[1], '  Blue  = Image token ({%token})', 'DDEEFF'),
        (hdr[2], '  Yellow = Boundary pair tokens', 'FFFACD'),
    ]:
        set_cell_bg(cell, color.strip())
        p = cell.paragraphs[0]
        p.add_run(txt).font.size = Pt(8)
    doc.add_paragraph()

    for template_name, mappings in all_data:
        # Section heading
        h = doc.add_heading(template_name, level=1)
        h.runs[0].font.size = Pt(13)
        h.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

        if not mappings:
            doc.add_paragraph('⚠️  No tokens matched in this template.').style.font.size = Pt(9)
            doc.add_paragraph()
            continue

        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = 'Table Grid'
        tbl.autofit = False

        # Set column widths
        widths = [DocxCm(7), DocxCm(6), DocxCm(4)]
        for i, w in enumerate(widths):
            for cell in tbl.columns[i].cells:
                cell.width = w

        # Header row
        hdr_cells = tbl.rows[0].cells
        for cell, label in zip(hdr_cells, ['Field Label (as in template)', 'Token Assigned', 'Notes']):
            set_cell_bg(cell, '1A56DB')
            p = cell.paragraphs[0]
            run = p.add_run(label)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Data rows
        for label, token, note in mappings:
            row = tbl.add_row()
            cells = row.cells

            is_image = token.startswith('{%')
            is_boundary = 'Boundary' in note

            bg = 'DDEEFF' if is_image else ('FFFACD' if is_boundary else 'FFFFFF')
            for c in cells:
                set_cell_bg(c, bg)

            cells[0].paragraphs[0].add_run(label).font.size = Pt(8)
            t_run = cells[1].paragraphs[0].add_run(token)
            t_run.font.size = Pt(8)
            t_run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B) if not is_image else RGBColor(0x15, 0x65, 0xC0)
            cells[2].paragraphs[0].add_run(note).font.size = Pt(8)

        doc.add_paragraph(f'Total tokens: {len(mappings)}').style.font.size = Pt(8)
        doc.add_paragraph()  # spacer between templates

    doc.save(OUTPUT)
    print(f'\n✓ Reference document saved to: {OUTPUT}')

# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    # Use backup if available (originals), else use templates dir
    source_dir = BACKUP_DIR if BACKUP_DIR.exists() and any(BACKUP_DIR.glob('*.docx')) else TEMPLATES_DIR
    print(f'Reading templates from: {source_dir}')

    templates = sorted(source_dir.glob('*.docx'))
    if not templates:
        print('ERROR: No .docx files found'); return

    all_data = []
    for p in templates:
        print(f'  Extracting: {p.name}')
        mappings = extract_mappings(p)
        all_data.append((p.stem, mappings))
        print(f'    → {len(mappings)} tokens')

    build_doc(all_data)

if __name__ == '__main__':
    main()
